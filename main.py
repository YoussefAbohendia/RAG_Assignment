import os
from pathlib import Path
from typing import List, Dict
import fitz  # PyMuPDF for PDFs
import docx  # python-docx for Word documents
from sklearn.metrics.pairwise import cosine_similarity
from langchain.text_splitter import CharacterTextSplitter
from sentence_transformers import SentenceTransformer
from langchain.embeddings import HuggingFaceEmbeddings
import faiss
import numpy as np
import pickle
from langchain.vectorstores import FAISS
from langchain.schema import Document
from transformers import pipeline

# ------------------------------
# Task 1.3 - Load documents
# ------------------------------
def load_documents(directory: str) -> List[Dict]:
    documents = []
    for file_path in Path(directory).glob("*"):
        try:
            if file_path.suffix.lower() == ".pdf":
                text = ""
                with fitz.open(file_path) as doc:
                    for page in doc:
                        text += page.get_text()
            elif file_path.suffix.lower() == ".docx":
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif file_path.suffix.lower() == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            else:
                print(f"⚠️ Unsupported file type: {file_path.name}")
                continue

            documents.append({
                "text": text,
                "metadata": {
                    "filename": file_path.name,
                    "path": str(file_path.resolve())
                }
            })

        except Exception as e:
            print(f"❌ Error reading {file_path.name}: {e}")
    return documents

# ------------------------------
# Task 2.1 - Split into chunks
# ------------------------------
def split_documents(docs: List[Dict], chunk_size=500, chunk_overlap=50) -> List[Dict]:
    splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len
    )

    chunks = []
    for doc in docs:
        text_chunks = splitter.split_text(doc["text"])
        for i, chunk in enumerate(text_chunks):
            chunks.append({
                "text": chunk,
                "metadata": {
                    **doc["metadata"],
                    "chunk": i + 1,
                    "text": chunk
                }
            })
    return chunks

# ------------------------------
# Task 2.2 - Embed and save to FAISS
# ------------------------------
def embed_and_store(chunks: List[Dict], model_name: str, faiss_index_path: str):
    print(f"\n🔍 Loading embedding model: {model_name}")
    model = SentenceTransformer(model_name)

    texts = [chunk["text"] for chunk in chunks]
    metadatas = [chunk["metadata"] for chunk in chunks]

    print("🧠 Generating embeddings...")
    embeddings = model.encode(texts, show_progress_bar=True)

    print("💾 Building FAISS index...")
    dim = embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(np.array(embeddings))

    faiss.write_index(index, faiss_index_path + ".index")
    with open(faiss_index_path + "_metadata.pkl", "wb") as f:
        pickle.dump(metadatas, f)

    print(f"✅ FAISS index saved to: {faiss_index_path}.index")
    print(f"✅ Metadata saved to: {faiss_index_path}_metadata.pkl")

# ------------------------------
# Task 3.1 - Load and search
# ------------------------------
def load_faiss_index(index_path: str):
    print("📥 Loading FAISS index and metadata...")
    index = faiss.read_index(index_path + ".index")
    with open(index_path + "_metadata.pkl", "rb") as f:
        metadata = pickle.load(f)
    return index, metadata

def search_documents(query: str, model_name: str, index_path: str, k=3):
    model = SentenceTransformer(model_name)
    query_embedding = model.encode([query])
    index, metadata = load_faiss_index(index_path)
    distances, indices = index.search(np.array(query_embedding), k)

    print(f"\n🔍 Top {k} Results for Query: \"{query}\"")
    for i, idx in enumerate(indices[0]):
        print(f"\nResult #{i+1} (Distance: {distances[0][i]:.4f})")
        print(f"📄 File: {metadata[idx]['filename']} | Chunk: {metadata[idx]['chunk']}")
        print("📝 Content Preview:")
        print("-" * 40)
        print(metadata[idx]['text'][:500] + "...\n")

# ------------------------------
# Task 3.2 - MMR Search
# ------------------------------
def mmr(query_embedding, doc_embeddings, k=3, lambda_param=0.5):
    selected = []
    unselected = list(range(len(doc_embeddings)))
    query_embedding = np.array(query_embedding).reshape(1, -1)
    doc_embeddings = np.array(doc_embeddings)

    sim_to_query = cosine_similarity(query_embedding, doc_embeddings)[0]
    sim_between_docs = cosine_similarity(doc_embeddings)

    for _ in range(k):
        mmr_scores = []
        for idx in unselected:
            diversity = 0 if not selected else max(sim_between_docs[idx][j] for j in selected)
            mmr_score = lambda_param * sim_to_query[idx] - (1 - lambda_param) * diversity
            mmr_scores.append((mmr_score, idx))

        mmr_scores.sort(reverse=True)
        best = mmr_scores[0][1]
        selected.append(best)
        unselected.remove(best)

    return selected

def search_documents_mmr(query: str, model_name: str, index_path: str, k=3):
    model = SentenceTransformer(model_name)
    query_embedding = model.encode([query])
    index, metadata = load_faiss_index(index_path)
    all_embeddings = np.array([index.reconstruct(i) for i in range(index.ntotal)])
    selected_indices = mmr(query_embedding, all_embeddings, k=k)

    print(f"\n🧠 MMR Top {k} Diverse Results for Query: \"{query}\"")
    for rank, idx in enumerate(selected_indices):
        print(f"\nResult #{rank+1}")
        print(f"📄 File: {metadata[idx]['filename']} | Chunk: {metadata[idx]['chunk']}")
        print("📝 Content Preview:")
        print("-" * 40)
        print(metadata[idx]['text'][:500] + "...\n")

# ------------------------------
# Task 4 - RAG Answer with HuggingFace Transformers
# ------------------------------
def generate_answer_with_huggingface(query: str, index_path: str, model_name="google/flan-t5-base"):
    print("\n🧠 Loading FAISS retriever and generating answer (HuggingFace)...")
    index = faiss.read_index(index_path + ".index")
    with open(index_path + "_metadata.pkl", "rb") as f:
        metadata = pickle.load(f)

    docs = [Document(page_content=meta["text"], metadata=meta) for meta in metadata]
    embedding_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    db = FAISS.from_documents(docs, embedding_model)
    retriever = db.as_retriever(search_type="similarity", search_kwargs={"k": 3})

    retrieved_docs = retriever.get_relevant_documents(query)
    context = "\n".join([doc.page_content for doc in retrieved_docs])

    prompt = f"""You are a helpful assistant. Use the context below to answer the question.

Context:
{context}

Question:
{query}

Answer:"""

    generator = pipeline("text2text-generation", model=model_name, tokenizer=model_name)
    response = generator(prompt, max_new_tokens=300)[0]['generated_text']

    print("\n💬 Answer:")
    print("-" * 50)
    print(response.strip())

# ------------------------------
# Main controller
# ------------------------------
if __name__ == "__main__":
    mode = input("Select mode: [1] Build Index  [2] Search  [3] MMR Search  [4] RAG Answer: ").strip()

    if mode == "1":
        docs = load_documents("documents")
        print(f"\n✅ Loaded {len(docs)} document(s).")
        chunks = split_documents(docs, chunk_size=500, chunk_overlap=50)
        print(f"🧩 Split into {len(chunks)} chunks.")
        os.makedirs("faiss_store", exist_ok=True)
        embed_and_store(
            chunks=chunks,
            model_name="all-MiniLM-L6-v2",
            faiss_index_path="faiss_store/index_miniLM"
        )

    elif mode == "2":
        query = input("\n💬 Enter your query: ")
        search_documents(
            query=query,
            model_name="all-MiniLM-L6-v2",
            index_path="faiss_store/index_miniLM",
            k=3
        )

    elif mode == "3":
        query = input("\n💬 Enter your query for MMR search: ")
        search_documents_mmr(
            query=query,
            model_name="all-MiniLM-L6-v2",
            index_path="faiss_store/index_miniLM",
            k=3
        )

    elif mode == "4":
        query = input("\n💬 Enter your question for RAG LLM answer: ")
        generate_answer_with_huggingface(
            query=query,
            index_path="faiss_store/index_miniLM",
            model_name="google/flan-t5-base"
        )

    else:
        print("❌ Invalid option. Choose 1, 2, 3, or 4.")
